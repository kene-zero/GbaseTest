from docx.oxml.ns import qn
from openpyxl import load_workbook
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls


def generate_testResult():
    # 读取 Excel 文件
    workbook = load_workbook('D:\\File\\测试文件\\测试用例\\测试用例.xlsx')  # 替换为你的 Excel 文件名
    sheet = workbook.active

    # 创建 Word 文档
    doc = Document()
    doc.styles['Normal'].font.name = "Times New Roman"
    doc.styles['Normal'].element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    title_list = ["测试说明", "前置条件", '测试步骤', '预期结果', '测试结果']
    pre = """1.GBase8C数据库状态正常
    2.GBase8A数据库状态正常
    3.Kafka集群以及Flink集群正常
    4.SCT工具以部署完成且状态正常
    """
    for num in range(2, 45):
        sheet_list = ["验证" + sheet[f'E{num}'].value, pre, sheet[f'H{num}'].value.replace('\n', ''),
                      sheet[f'I{num}'].value.replace('\n', ''), "通过"]
        table = doc.add_table(rows=5, cols=7)
        table.cell(0, 1).merge(table.cell(0, 6))
        table.cell(1, 1).merge(table.cell(1, 6))
        table.cell(2, 1).merge(table.cell(2, 6))
        table.cell(3, 1).merge(table.cell(3, 6))
        table.cell(4, 1).merge(table.cell(4, 6))
        table.style = 'Table Grid'

        for x in range(0, 5):
            table.cell(x, 0).text = title_list[x]
            clo = parse_xml(r'<w:shd {} w:fill="#D3D3D3"/>'.format(nsdecls('w')))
            table.cell(x, 0)._tc.get_or_add_tcPr().append(clo)
            table.cell(x, 0).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in table.cell(x, 0).paragraphs:
                for run in paragraph.runs:
                    run.bold = True

        for step in range(0, 5):
            table.cell(step, 1).text = sheet_list[step].strip()
            for paragraph in table.cell(step, 1).paragraphs:
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                for run in paragraph.runs:
                    run.font.name = '宋体'

        doc.add_paragraph('')

    # 保存 Word 文档
    doc.save('D:\\File\\测试文件\\测试用例\\excel_data_in_word.docx')  # 替换为你想要保存的 Word 文档文件名
